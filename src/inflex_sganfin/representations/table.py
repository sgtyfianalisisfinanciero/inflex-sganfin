import itertools
from pathlib import Path

import docx
import docx.document
import docx.enum
import docx.enum.table
import docx.enum.text
import docx.oxml
import docx.shared
import docx.table
from data.series import Series
from docx.oxml.ns import qn


class TableColumn:
    def __init__(self, name: str) -> None:
        self.name: str = name
        # loc: SeriesLocation = None
        # units: Unit = None
        # precision: int = 0
        # color: bool = True
        # highlight_outliers: bool = True
        # positive_color: str = ""
        # negative_color: str = ""


class TableBlock:
    def __init__(
        self, heading: str, columns: list[TableColumn], rows: list[Series]
    ) -> None:
        self.heading: str = heading
        self.columns: list[TableColumn] = columns
        self.rows: list[Series] = rows

    @property
    def column_names(self) -> list[str]:
        return [column.name for column in self.columns]


class Table:

    def __init__(
        self,
        blocks: list[TableBlock],
        block_separation: bool = False,
        stack_axis: int = 0,
    ) -> None:
        self.blocks: list[TableBlock] = blocks
        self.block_separation: bool = block_separation
        self.stack_axis: int = stack_axis

    def render_docx(self, docx_document: docx.document.Document):
        self.horizontal: bool = self.stack_axis == 1
        docx_table: docx.table.Table = self._create_docx_table(docx_document)
        if self.horizontal:
            # if the table is vstacked there are no headers to render
            self._render_docx_hstacked_table_header(docx_table)
        self._render_docx_column_names(docx_table)
        self._render_docx_series_names(docx_table)
        self._render_docx_content(docx_table)
        docx_document.save(Path("test.docx"))

    def _create_docx_table(
        self, docx_document: docx.document.Document
    ) -> docx.table.Table:
        block_rows: list[int] = [len(block.rows) for block in self.blocks]
        block_columns: list[int] = [len(block.columns) for block in self.blocks]
        if self.horizontal:
            # two extra rows, one for block names and one for column names
            render_rows = 2 + max(block_rows)
            # one extra column for series names
            render_columns = 1 + sum(block_columns)
        else:
            # only one extra row for column names
            render_rows = 1 + sum(block_rows)
            # one extra column for series names
            render_columns = 1 + max(block_columns)
        docx_table: docx.table.Table = docx_document.add_table(
            rows=render_rows, cols=render_columns
        )
        return docx_table

    def _render_docx_style(self, docx_table: docx.table.Table) -> None:
        # TODO: remove hardcoded parameters
        docx_table.style = "Light Shading Accent 1"
        docx_table.autofit = False

    def _render_docx_hstacked_table_header(self, docx_table: docx.table.Table) -> None:
        counter = 1  # leave the first column for series names
        for block in self.blocks:
            # create a wide enough cell
            block_columns: int = len(block.columns)
            cell: docx.table._Cell = docx_table.cell(0, counter)
            for _ in range(block_columns - 1):
                counter = counter + 1
                cell.merge(docx_table.cell(0, counter))
            # fill the wide cell
            cell.text = block.heading
            # TODO: replace hardcoded parameters
            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(12)
            counter = counter + 1

    @property
    def column_names(self) -> list[str]:
        column_names_nested: list[list[str]] = [
            block.column_names for block in self.blocks
        ]
        column_names: list[str] = list(
            itertools.chain.from_iterable(column_names_nested)
        )
        return column_names

    def _render_docx_column_names(self, docx_table: docx.table.Table) -> None:
        # the first row is reserved for block header names in horizontal tables
        start_row: int = 1 if self.horizontal else 0
        for column, name in enumerate(self.column_names):
            # leave the first column free for series names
            target_cell: docx.table._Cell = docx_table.cell(start_row, column + 1)
            target_cell.text = name
            # TODO: replace hardcoded parameters
            target_cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(10)
            target_cell.paragraphs[0].alignment = (
                docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            )
            target_cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    def _render_docx_series_names(self, docx_table: docx.table.Table) -> None:
        # One extra row for hstacked tables for block header names
        start_row: int = 2 if self.horizontal else 1
        filled_rows: int = 0
        for block in self.blocks:
            for series in block.rows:
                target_cell: docx.table._Cell = docx_table.cell(
                    start_row + filled_rows, 0
                )
                target_cell.text = series.report_name
                # TODO: replace hardcoded parameters
                target_cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(9)
                target_cell.paragraphs[0].runs[0].font.bold = True
                target_cell.width = docx.shared.Inches(1)
                filled_rows = filled_rows + 1
            if self.block_separation and not self.horizontal:
                # cell separation
                # TODO: replace hardcoded parameters
                self._separate_cell(
                    target_cell,
                    bottom={"sz": 1, "val": "double", "color": "#000000", "space": 2},
                )
            if self.horizontal:
                # Series report names must be the same for all blocks
                break

    def _separate_cell(self, cell: docx.table._Cell, **kwargs) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = docx.oxml.OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)

        for edge in ["start", "top", "end", "bottom", "insideH", "insideV"]:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f"w:{edge}"
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = docx.oxml.OxmlElement(tag)
                    tcBorders.append(element)
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn(f"w:{key}"), str(edge_data[key]))

    def _render_docx_content(self, docx_table: docx.table.Table) -> None:
        start_row: int = 2 if self.horizontal else 1
        filled_rows: int = 0
        filled_cols: int = 0
        for block in self.blocks:
            for i in range(
                start_row + filled_rows, start_row + filled_rows + len(block.rows)
            ):
                for j in range(1 + filled_cols, 1 + filled_cols + len(block.columns)):
                    target_cell: docx.table._Cell = docx_table.cell(i, j)
                    target_cell.text = "NA"
                    # TODO: replace hardcoded parameters
                    target_cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(9)
                    target_cell.paragraphs[0].alignment = (
                        docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    )
                    target_cell.vertical_alignment = (
                        docx.enum.table.WD_ALIGN_VERTICAL.CENTER
                    )
            if self.block_separation and not self.horizontal:
                # cell separation
                # TODO: replace hardcoded parameters
                self._separate_cell(
                    target_cell,
                    bottom={"sz": 1, "val": "double", "color": "#000000", "space": 2},
                )

            if self.horizontal:
                filled_cols = filled_cols + len(block.columns)
            else:
                filled_rows = filled_rows + len(block.rows)
